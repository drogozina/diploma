# -*- coding: utf-8 -*-
from django.http import HttpResponse
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Inches, Cm

def getDate(date):
    array = date.split('-')
    string = array[2]+'.'+array[1]+'.'+array[0]
    return string

def getArray(string):
    try:
        return string.split(',')
    except:
        return string

def getInfoFromDB(string):
    array = string.split('$')
    return array

def make(program):
    document = Document()
    styles = document.styles
    sections = document.sections
    section = sections[0]
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    insidelist = styles['List Number']
    insidelist.font.name = 'Times New Roman'
    insidelist.font.size = Pt(12)
    insidelist.font.bold = True

    # style for doc's hat
    hat = styles.add_style('Hat', WD_STYLE_TYPE.PARAGRAPH)
    font = hat.font
    hat.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font.name = 'Times New Roman'
    font.size = Pt(12)
    ###
    ### МИНОБРНАУКИ РОССИИ
    minobrnayki = document.add_paragraph(
        program.document.minobrnayki + '\n' + program.document.fed)
    minobrnayki.style = hat
    ####

    omsustyle = styles.add_style('OMSU', WD_STYLE_TYPE.PARAGRAPH)
    font = omsustyle.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    omsustyle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omsu = document.add_paragraph('«Омский государственный университет им. Ф.М. Достоевского»')
    omsu.style = omsustyle

    kafedra = document.add_paragraph('Факультет компьютерных наук\n' + program.dep)
    kafedra.style = hat

    leftpar = styles.add_style('LeftPar', WD_STYLE_TYPE.PARAGRAPH)
    paragr_format = leftpar.paragraph_format
    paragr_format.left_indent = Cm(10)
    paragr_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    leftpar.font.name = 'Times New Roman'
    leftpar.font.size = Pt(12)
    confirm = document.add_paragraph('\n\n«Утверждаю»\n'
                                     +program.document.approvePosition+'\n'
                                     '_______________'+program.document.approve+'\n'
                                     '«_____» _____________ 201__ г.\n\n\n\n')
    confirm.style = leftpar

    rabprogr = document.add_paragraph('Рабочая программа дисциплины\n')
    rabprogr.style = omsustyle

    # DISC.NAME
    discipname = document.add_paragraph('«' + program.discipline.name + '»\n\n\n')

    disp = styles.add_style('DISCIPLINE', WD_STYLE_TYPE.PARAGRAPH)
    font = disp.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    font.italic = True
    disp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    discipname.style = disp

    table = document.add_table(rows=5, cols=2, style='Table Grid')
    # tablestyle = styles.add_style('Table Grid', WD_STYLE_TYPE.TABLE)
    table.style.font.name = 'Times New Roman'
    table.style.font.size = Pt(12)
    # table.style = tablestyle
    table.rows[0].cells[0].text = 'Направление подготовки'
    table.rows[1].cells[0].text = 'Код направления подготовки'
    table.rows[2].cells[0].text = 'Профиль подготовки'
    table.rows[3].cells[0].text = 'Базовая/вариативная часть'
    table.rows[4].cells[0].text = 'Обязательная дисциплина/дисциплина по выбору/факультатив'
    table.rows[0].cells[1].text = program.group.dirname
    table.rows[1].cells[1].text = program.group.direction
    table.rows[2].cells[1].text = program.group.profile
    table.rows[3].cells[1].text = program.discipline.part
    table.rows[4].cells[1].text = program.discipline.kindofdis
    import datetime
    now = datetime.datetime.now()
    document.add_paragraph('\n\n\n\n\n\n\n\n\nг. Омск – ' + str(now.year) + ' г.', style=hat)
    document.add_page_break()

    # PAGE 2
    page2 = styles.add_style('Page2', WD_STYLE_TYPE.PARAGRAPH)
    font = page2.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    page2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(
        'Программа дисциплины «'+program.discipline.name+'» разработана:\n' + program.degree + '     ___________________     ' + program.author,
        style=page2)
    kafname = program.dep.split(' ')
    document.add_paragraph(
        'Программа рассмотрена на заседании кафедры ' + kafname[1] + ' (протокол № ' + program.protocolnum + ' от ' + getDate(program.date) + ') ',
        style=page2)
    document.add_paragraph(
        'Программа разработана в соответствии с федеральным государственным образовательным стандартом высшего образования РФ и согласована с факультетом, реализующим направление ' + program.group.direction + ' ' + program.group.dirname,
        style=page2)
    document.add_paragraph('Декан факультета компьютерных наук		 	___________________     '+program.document.decan,
                           style=page2)
    document.add_paragraph(
        'Зав. библиотекой _____________________________    ________________________________									(подпись, Ф.И.О. проверяющего)											М.П.',
        style=page2)
    document.add_page_break()

    # PAGE 3
    place = document.add_paragraph('Место дисциплины в структуре ОПОП.', 'ListNumber')

    page3 = styles.add_style('Page3', WD_STYLE_TYPE.PARAGRAPH)
    font = page3.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    page3.paragraph_format.first_line_indent = Cm(1)
    page3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph(
        'Цель дисциплины: ' + program.aim,'Page3')
    document.add_paragraph('Задачи дисциплины: ' + program.goal, 'Page3')
    document.add_paragraph(
        'Перечень планируемых результатов обучения по дисциплине, соотнесенных с планируемыми результатами освоения образовательной программы',
        "OMSU")

    planResults = document.add_table(rows=3, cols=3, style='Table Grid')
    planResults.style.font.name = 'Times New Roman'
    planResults.style.font.size = Pt(12)

    planResults.rows[0].cells[0].text = 'Коды ком-петенций (по ФГОС ВО)'
    planResults.rows[0].cells[1].text = 'Содержание компетен-ций согласно ФГОС ВО'
    planResults.rows[0].cells[2].text = 'Перечень планируемых результатов обучения по дисци-плине'
    if "ПСК" in program.discipline.competention:
        planResults.rows[1].cells[0].text = 'ПСК'
        planResults.rows[1].cells[1].text = 'ПРОФИЛЬНО-СПЕЦИАЛИЗИРОВАННЫЕ КОМПЕТЕНЦИИ ВЫПУСКНИКА'
    else:
        planResults.rows[1].cells[0].text = 'ОК'
        planResults.rows[1].cells[1].text = 'ОБЩЕКУЛЬТУРНЫЕ КОМПЕТЕНЦИИ ВЫПУСКНИКА'
    planResults.rows[1].cells[1].merge(planResults.rows[1].cells[2])

    array = getInfoFromDB(program.planResults)
    planResults.rows[2].cells[0].text = program.discipline.competention
    planResults.rows[2].cells[1].text = array[0]
    planResults.rows[2].cells[2].text = array[1]

    document.add_paragraph(
        'Содержание дисциплины, структурированное по разделам, с указанием отведенного на них количества академических часов и видов учебных занятий',
        'ListNumber')

    tableWithHours = document.add_table(rows=4, cols=8, style='Table Grid')
    tableWithHours.style.font.name = 'Times New Roman'
    tableWithHours.style.font.size = Pt(12)
    for i in range(1, 4):
        tableWithHours.rows[0].cells[0].merge(tableWithHours.rows[i].cells[0])
        tableWithHours.rows[0].cells[1].merge(tableWithHours.rows[i].cells[1])
        tableWithHours.rows[0].cells[2].merge(tableWithHours.rows[i].cells[2])
    for i in range(2, 4):
        tableWithHours.rows[1].cells[7].merge(tableWithHours.rows[i].cells[7])
    tableWithHours.rows[2].cells[3].merge(tableWithHours.rows[3].cells[3])
    for i in range(4, 8):
        tableWithHours.rows[0].cells[3].merge(tableWithHours.rows[0].cells[i])
    for i in range(4, 7):
        tableWithHours.rows[1].cells[3].merge(tableWithHours.rows[1].cells[i])
        tableWithHours.rows[2].cells[4].merge(tableWithHours.rows[2].cells[i])
    tableWithHours.rows[0].cells[0].text = 'Раздел дисциплины'
    tableWithHours.rows[0].cells[1].text = 'Семестр'
    tableWithHours.rows[0].cells[2].text = 'Всего по разделу'
    tableWithHours.rows[0].cells[3].text = 'Объем часов (по видам работы)'
    tableWithHours.rows[1].cells[3].text = 'Контактная работа'
    tableWithHours.rows[1].cells[7].text = 'Самостоятельная работа обучающихся'
    tableWithHours.rows[2].cells[3].text = 'Всего'
    tableWithHours.rows[2].cells[4].text = 'в том числе'
    tableWithHours.rows[3].cells[4].text = 'Занятия лекционного типа'
    tableWithHours.rows[3].cells[5].text = 'Практические и семинарские занятия'
    tableWithHours.rows[3].cells[6].text = 'Лабораторные занятия'
    array = getArray(program.disciplineContent)
    hour = ''
    for item in array:
        new = getInfoFromDB(item)
        cells = tableWithHours.add_row().cells
        i = 0
        for cell in cells:
            try:
                cell.text = new[i]
                if i == 2:
                    hour = new[i]
            except:
                pass
            i += 1

    if program.intensity[len(program.intensity)-1] == '1':
        zachet = ' зачетная единица; '
    else:
        if program.intensity[0] == '1' and len(program.intensity)==2:
            zachet = ' зачетных единиц; '
        else:
            if program.intensity[len(program.intensity)-1] == '2' or program.intensity[len(program.intensity)-1] == '3' or program.intensity[len(program.intensity)-1] == '4':
                zachet = ' зачетные единицы; '
            else:
                zachet = ' зачетных единиц; '
    if hour[len(hour)-1] == '1':
        kolvohours = ' час.'
    else:
        if hour[0] == '1' and len(hour) == 2:
            kolvohours = ' часов.'
        else:
            if hour[len(hour)-1] == '2' or hour[len(hour)-1] == '3' or hour[len(hour)-1] == '4':
                kolvohours = ' часа.'
            else:
                kolvohours = ' часов.'
    document.add_paragraph('Общая трудоемкость дисциплины: ' + program.intensity + zachet + hour + kolvohours, 'Page2')
    document.add_paragraph('Форма промежуточной аттестации: ' + program.attestation, 'Page2')

    document.add_paragraph('Содержание дисциплины', 'ListNumber')

    insidelist = styles['List Number 2']
    insidelist.font.name = 'Times New Roman'
    insidelist.font.size = Pt(12)
    insidelist.paragraph_format.left_indent = Cm(0.63)
    insidelist.paragraph_format.first_line_indent = Cm(0.85)
    insidelist.font.bold = True
    insidelist = styles['List Continue 2']
    insidelist.font.name = 'Times New Roman'
    insidelist.font.size = Pt(12)
    insidelist.paragraph_format.left_indent = Cm(0.63)
    insidelist.paragraph_format.first_line_indent = Cm(0.85)

    document.add_paragraph('Лекции', 'List Number 2')
    document.add_paragraph('Темы лекций, их содержание, трудоемкость', 'List Continue 2')
    document.add_paragraph(program.lectures, 'List Continue 2')

    document.add_paragraph('Практические занятия', 'List Number 2')
    document.add_paragraph('Состав и объем практических занятий', 'List Continue 2')
    document.add_paragraph(program.pracs, 'List Continue 2')

    document.add_paragraph('Лабораторные занятия', 'List Number 2')
    document.add_paragraph('Состав и объем лаборатоных занятий', 'List Continue 2')
    document.add_paragraph(program.labs, 'List Continue 2')

    document.add_paragraph(
        'Перечень учебно-методического обеспечения для самостоятельной работы обучающихся по дисциплине' , 'ListNumber')
    document.add_paragraph(program.listLit, 'List Continue 2')
    document.add_paragraph('Фонд оценочных средств для проведения промежуточной аттестации обучающихся по дисциплине ',
                           'ListNumber')
    document.add_paragraph(program.fond, 'List Continue 2')
    document.add_paragraph('Перечень основной и дополнительной литературы', 'ListNumber')
    document.add_paragraph('Основная литература', 'List Number 2')
    document.add_paragraph(program.mainlit, 'List Continue 2')
    document.add_paragraph('Дополнительная литература', 'List Number 2')
    document.add_paragraph(program.addlit, 'List Continue 2')
    document.add_paragraph('Перечень ресурсов информационно-телекоммуникационной сети Интернет', 'ListNumber')
    document.add_paragraph(program.internet, 'List Continue 2')
    document.add_paragraph('Методические указания для обучающихся по освоению дисциплины', 'ListNumber')
    document.add_paragraph(program.method, 'List Continue 2')

    document.add_paragraph('Методическое обеспечение самостоятельной работы обучающихся', "OMSU")
    methodTable = document.add_table(rows=1, cols=4, style='Table Grid')
    methodTable.rows[0].cells[0].text = 'Тема, раздел'
    methodTable.rows[0].cells[1].text = 'Задания для самостоятельной работы'
    methodTable.rows[0].cells[2].text = 'Список литературы (с указанием разделов, глав/страниц)'
    methodTable.rows[0].cells[3].text = 'Форма контроля'
    array = getArray(program.methodTable)

    for item in array:
        new = getInfoFromDB(item)
        cells = methodTable.add_row().cells
        i = 0
        for cell in cells:
            try:
                cell.text = new[i]
            except:
                pass
            i += 1

    document.add_paragraph('Подготовка к лекции заключается в следующем: \n' + program.preplec, 'Page2')
    document.add_paragraph('Подготовка к практическим занятиям включает: \n' + program.preprac, 'Page2')
    document.add_paragraph('Подготовка к экзамену: \n' + program.prepex, 'Page2')

    document.add_paragraph(
        'Перечень информационных технологий, используемых при осуществлении образовательного процесса по дисциплине, включая перечень программного обеспечения и информационных справочных систем',
        'ListNumber')
    document.add_paragraph(program.listit, 'List Continue 2')
    document.add_paragraph(
        'Описание материально-технической базы, необходимой для осуществления образовательного процесса по дисциплине',
        'ListNumber')

    document.add_paragraph(
        'Материально-техническая база для проведения всех видов дисциплинарной и междисциплинарной подготовки', "OMSU")
    methodTable = document.add_table(rows=1, cols=3, style='Table Grid')
    methodTable.rows[0].cells[0].text = 'Наименование специальных помещений и помещений для самостоятельной работы'
    methodTable.rows[0].cells[
        1].text = 'Оснащенность специальных помещений и помещений для самостоятельной работы обучающихся'
    methodTable.rows[0].cells[
        2].text = 'Перечень лицензионного программного обеспечения. Реквизиты подтверждающего документа'
    array = getArray(program.tableBase)
    for item in array:
        new = getInfoFromDB(item)
        cells = methodTable.add_row().cells
        i = 0
        for cell in cells:
            try:
                cell.text = new[i]
            except:
                pass
            i += 1
    document.save('C:/Users/even-/YandexDisk/DIPLOMA/templates/Program.docx')
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)
    return response
